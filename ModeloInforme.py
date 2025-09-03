#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from pathlib import Path
import re, ollama
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# ---------------- FUNCIONES DE FORMATO ---------------- #

def aplicar_fuente_run(run, fuente="Arial", tam=11, bold=False):
    run.font.size = Pt(tam)
    run.bold = bool(bold)
    try:
        run.font.name = fuente
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), fuente)
        rFonts.set(qn('w:hAnsi'), fuente)
    except:
        pass

def insertar_parrafo_despues(par_ref, txt, bold=False, tam=11):
    doc = par_ref._parent
    nuevo = doc.add_paragraph()
    par_ref._p.addnext(nuevo._p)
    run = nuevo.add_run(txt)
    aplicar_fuente_run(run, tam=tam, bold=bold)
    nuevo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return nuevo

def encontrar_parrafo_con_ancla(doc, anchor):
    return next((p for p in doc.paragraphs if anchor in p.text), None)

# ---------------- LIMPIEZA ---------------- #

def clean_text(txt):
    if not txt:
        return ""
    txt = txt.replace("“", '"').replace("”", '"').replace("’", "'")
    return txt.strip()

def extraer_contenido_ollama(resp):
    """
    Extrae el texto real de la respuesta de ollama.chat
    Maneja distintos formatos posibles.
    """
    try:
        if isinstance(resp, dict):
            m = resp.get("message")
            if isinstance(m, dict) and "content" in m:
                return clean_text(m["content"])
            if "content" in resp and isinstance(resp["content"], str):
                return clean_text(resp["content"])
        if hasattr(resp, "message"):
            m = resp.message
            if isinstance(m, dict) and "content" in m:
                return clean_text(m["content"])
            if hasattr(m, "content"):
                return clean_text(m.content)
        if hasattr(resp, "content"):
            return clean_text(resp.content)
    except Exception:
        pass
    return clean_text(str(resp))

def sanitize_no_periods(text):
    if not text:
        return text
    text = re.sub(r'\.\s*', ', ', text)
    text = text.replace('.', ',')
    text = re.sub(r',\s*,+', ', ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'[,\s]+$', '', text)
    return text

# ---------------- PROMPTS ---------------- #

def prompt_intro():
    return (
        "RESPONDE SOLO EN ESPAÑOL. DEVUELVE SÓLO EL TEXTO SOLICITADO.\n\n"
        "INTRODUCCION:\n"
        "La introducción debe ser UN ÚNICO PÁRRAFO corrido de 4 a 6 líneas, no uses el caracter punto '.' en esta sección, "
        "usa comas y punto y coma si es necesario, no uses viñetas, no uses numeración, no uses saltos de párrafo"
    )

def prompt_para_punto(titulo):
    return (
        "RESPONDE SOLO EN ESPAÑOL. GENERA EXACTAMENTE ESTE FORMATO:\n\n"
        "Descripción:\n"
        "Texto corrido de 5 a 7 líneas, sin puntos '.', usa comas y punto y coma\n\n"
        "Ejemplo:\n"
        "```python\n"
        "print('ejemplo')\n"
        "```\n\n"
        "Explicación:\n"
        "Texto corrido de 3 a 5 líneas, sin puntos '.', usa comas y punto y coma\n\n"
        f"Tema: {titulo}"
    )

# ---------------- GENERADORES ---------------- #

def generar_intro(reintentos=2):
    for _ in range(reintentos+1):
        r = ollama.chat(model="mistral", messages=[{"role":"user","content":prompt_intro()}])
        txt = extraer_contenido_ollama(r)
        if txt:
            t = re.sub(r'^(INTRODUCCION:?)\s*', '', txt, flags=re.I).strip()
            t = " ".join(t.split())
            return sanitize_no_periods(t)
    return "Introducción no disponible"

def generar_contenido(titulo, reintentos=2):
    for _ in range(reintentos+1):
        r = ollama.chat(model="mistral", messages=[{"role":"user","content":prompt_para_punto(titulo)}])
        txt = extraer_contenido_ollama(r)
        partes = parsear_partes(txt)
        partes["descripcion"] = sanitize_no_periods(partes["descripcion"])
        partes["explicacion"] = sanitize_no_periods(partes["explicacion"])
        if partes["descripcion"] and partes["ejemplo"] and partes["explicacion"]:
            return partes
    return {
        "descripcion": "Descripción no disponible",
        "ejemplo": "# Ejemplo no disponible",
        "explicacion": "Explicación no disponible"
    }

def parsear_partes(txt):
    partes = {"descripcion": "", "ejemplo": "", "explicacion": ""}
    if not txt:
        return partes
    s = txt.replace("\r\n", "\n")
    def extract(start_label, end_label=None):
        i = s.find(start_label)
        if i == -1:
            return ""
        start = i + len(start_label)
        end = s.find(end_label) if end_label and s.find(end_label) != -1 else len(s)
        return s[start:end].strip()
    partes["descripcion"] = " ".join(extract("Descripción:", "Ejemplo:").split())
    ej_block = extract("Ejemplo:", "Explicación:")
    m = re.search(r"```(?:python)?\n(.*?)\n```", ej_block, flags=re.DOTALL | re.IGNORECASE)
    partes["ejemplo"] = m.group(1).strip() if m else ej_block.strip()
    partes["explicacion"] = " ".join(extract("Explicación:", None).split())
    return partes

# ---------------- GENERADOR DE INFORME ---------------- #

def generar_informe(doc_base, salida, puntos, tarea, ancla="[[INICIO_INFORME]]"):
    doc = Document(str(doc_base))
    par_ancla = encontrar_parrafo_con_ancla(doc, ancla) or doc.paragraphs[-1]

    t1 = insertar_parrafo_despues(par_ancla, "Tarea más significativa:", bold=True, tam=11)
    t2 = insertar_parrafo_despues(t1, tarea, bold=True, tam=11)
    d1 = insertar_parrafo_despues(t2, "Descripción del proceso:", bold=True, tam=11)
    intro_title = insertar_parrafo_despues(d1, "INTRODUCCION", bold=True, tam=18)

    intro = generar_intro()
    p_intro = insertar_parrafo_despues(intro_title, intro, bold=False, tam=11)

    ultimo = p_intro
    for titulo in puntos:
        p_t = insertar_parrafo_despues(ultimo, titulo, bold=True, tam=18)
        partes = generar_contenido(titulo)

        p_desc = insertar_parrafo_despues(p_t, partes["descripcion"], bold=False, tam=11)
        p_ej_label = insertar_parrafo_despues(p_desc, "Ejemplo:", bold=True, tam=11)
        p_code = insertar_parrafo_despues(p_ej_label, partes["ejemplo"], bold=False, tam=11)
        p_ex_label = insertar_parrafo_despues(p_code, "Explicación:", bold=True, tam=11)
        p_ex = insertar_parrafo_despues(p_ex_label, partes["explicacion"], bold=False, tam=11)

        ultimo = insertar_parrafo_despues(p_ex, "", bold=False, tam=11)

    doc.save(str(salida))
    return salida
