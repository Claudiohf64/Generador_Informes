# ModeloCuadro.py
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import ollama  # Cliente local de Ollama

EMU_PER_INCH = 914400
TWIPS_PER_CM = 567


# --- Funciones auxiliares ---
def set_row_height(row, height_cm=None, rule="atLeast"):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    if height_cm is not None:
        trHeight.set(qn("w:val"), str(int(height_cm * TWIPS_PER_CM)))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def set_col_widths(tbl, widths_in):
    # Intenta aplicar el ancho a cada celda por fila (compatible con python-docx)
    for row in tbl.rows:
        for j, w in enumerate(widths_in):
            try:
                row.cells[j].width = Inches(w)
            except Exception:
                # no detener ejecución si falla
                pass


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    cellMar = OxmlElement("w:tcMar")
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = OxmlElement(f"w:{margin_name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        cellMar.append(node)
    tcPr.append(cellMar)


# --- Normalización y parseo de horas ---
def _normalize_dash(text):
    if not text:
        return text
    return text.replace("—", "–").replace("-", "–").replace("– ", "–").replace(" –", "–").strip()


def parse_hora_rango_a_minutos(rango):
    """
    Acepta rangos con guion '–' y admite formatos:
      - '10:00 AM – 12:15 PM' (12h)
      - '07:30 – 17:00'         (24h)
      - también admite sin espacio '10:00AM–12:15PM'
    Retorna minutos totales del intervalo (si end < start, asume cruce de medianoche).
    """
    try:
        if not rango or "–" not in rango:
            return 0
        rango = _normalize_dash(rango)
        start_raw, end_raw = rango.split("–", 1)
        start = start_raw.strip()
        end = end_raw.strip()

        fmts = ("%I:%M %p", "%I:%M%p", "%H:%M", "%H:%M:%S")
        h1 = h2 = None
        for fmt in fmts:
            try:
                h1 = datetime.strptime(start, fmt)
                break
            except Exception:
                continue
        for fmt in fmts:
            try:
                h2 = datetime.strptime(end, fmt)
                break
            except Exception:
                continue

        # Si no se pudo parsear alguno, retorna 0
        if h1 is None or h2 is None:
            return 0

        # Comparar tiempos solo por hora:minuto (usar fecha fija)
        base = datetime(2000, 1, 1)
        dt1 = base.replace(hour=h1.hour, minute=h1.minute, second=h1.second)
        dt2 = base.replace(hour=h2.hour, minute=h2.minute, second=h2.second)

        diff = (dt2 - dt1).total_seconds() // 60
        if diff < 0:
            # cruzó medianoche: sumar 24h
            diff = (dt2 + timedelta(days=1) - dt1).total_seconds() // 60
        return int(diff)
    except Exception:
        return 0


def minutos_a_horas_minutos_str(minutos):
    h = minutos // 60
    m = minutos % 60
    if m == 0:
        return f"{h}H"
    else:
        return f"{h}H {m}M"


# --- Generador de descripciones con Mistral (ollama) ---
def generar_descripcion_tarea_mistral(tema, tarea):
    """
    Llama a Mistral vía ollama y garantiza:
      - respuesta en español
      - empieza con 'Realicé' (forzado si hace falta)
      - una sola oración corta (fallback si falla)
    """
    prompt = f"""
Eres un estudiante redactando un informe semanal en español.

Tema: {tema}
Tarea: {tarea}

Instrucciones estrictas:
- Responde SOLO en español
- Escribe exactamente UNA oración en primera persona y en pasado
- La oración DEBE comenzar con "Realicé"
- Usa máximo 15 palabras
- NO uses punto aparte ni saltos de línea
- NO uses viñetas ni numeración
- NO expliques de más
- NO uses signos de puntuación excepto comas

Ejemplo: "Realicé ejercicios de álgebra lineal usando matrices"
"""
    try:
        resp = ollama.chat(model="mistral", messages=[{"role": "user", "content": prompt}])
        text = resp.get("message", {}).get("content", "").strip()
        # limpiar saltos y puntos sobrantes
        text = text.replace("\n", " ").replace("..", " ").replace(".", "").strip()
        if not text:
            raise ValueError("Respuesta vacía")
        if not text.lower().startswith("realicé"):
            text = "Realicé " + text.lstrip()
        # limitar longitud a ~15 palabras (por si Mistral se va largo)
        palabras = text.split()
        if len(palabras) > 20:
            text = " ".join(palabras[:20])
        return text
    except Exception:
        # Fallback seguro y corto (no depende de la IA)
        t_short = tarea.strip()
        if len(t_short) > 60:
            t_short = t_short[:57].rsplit(" ", 1)[0] + "..."
        fallback = f"Realicé {t_short}"
        return fallback


def generar_descripciones_mistral(tema, tareas):
    """
    Genera lista de dicts {'nombre': tarea, 'descripcion': desc}
    Si una tarea falla individualmente, se aplica fallback para esa tarea.
    """
    resultados = []
    for t in tareas:
        try:
            desc = generar_descripcion_tarea_mistral(tema, t) if t else ""
        except Exception:
            desc = f"Realicé {t[:60]}" if t else "Realicé una tarea sin descripción"
        resultados.append({"nombre": t, "descripcion": desc})
    return resultados


# --- Construcción de la tabla en el documento ---
def build_table_in_doc(doc, dias_semana, table_style="Table Grid", day_w=1.10, hours_w=1.10, data_row_height_cm=2.8):
    sec = doc.sections[0]
    usable_in = (sec.page_width - sec.left_margin - sec.right_margin) / EMU_PER_INCH
    acts_w = usable_in - day_w - hours_w

    filas = len(dias_semana) + 2  # header + total
    tabla = doc.add_table(rows=filas, cols=3)
    tabla.style = table_style
    tabla.autofit = False
    set_col_widths(tabla, [day_w, acts_w, hours_w])

    # Márgenes internos
    for row in tabla.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=100, start=100, bottom=100, end=100)

    # Header
    hdr = tabla.rows[0].cells
    hdr[0].text = "DÍA"
    hdr[1].text = "ACTIVIDADES/TRABAJOS EFECTUADOS"
    hdr[2].text = "HORAS"
    for c in hdr:
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(tabla.rows[0], height_cm=1.0, rule="atLeast")

    total_minutos = 0
    for i, dia_info in enumerate(dias_semana, start=1):
        c_dia, c_act, c_horas = tabla.rows[i].cells

        # Día y fecha
        p_dia = c_dia.paragraphs[0]
        fecha_text = dia_info.get("fecha", "")
        p_dia.text = f"{dia_info.get('dia', '')} {fecha_text}".strip()
        p_dia.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_dia.runs:
            run.font.size = Pt(9)
            run.bold = False
        c_dia.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Limpiar contenido previo en actividades (manteniendo estructura)
        try:
            for para in list(c_act.paragraphs):
                p_element = para._element
                p_element.getparent().remove(p_element)
            # Asegurar limpieza completa como en código antiguo
            try:
                c_act._element.clear_content()
            except Exception:
                # si no existe clear_content, ignorar (ya removimos párrafos)
                pass
        except Exception:
            pass

        if not dia_info.get("laborable", False):
            p = c_act.add_paragraph(dia_info.get("razon_no_lab", "Día no laborable"))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(7)
                run.bold = False
        else:
            for tema_info in dia_info.get("temas", []):
                tema_text = tema_info.get("tema", "Sin tema")
                p_tema = c_act.add_paragraph(tema_text)
                p_tema.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p_tema.runs:
                    run.font.size = Pt(9)
                    run.bold = True

                for tarea in tema_info.get("tareas", []):
                    nombre = tarea.get("nombre", "")
                    descripcion = tarea.get("descripcion", "")

                    p_nom = c_act.add_paragraph(nombre)
                    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p_nom.runs:
                        run.font.size = Pt(7)
                        run.bold = True

                    p_desc = c_act.add_paragraph(descripcion)
                    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p_desc.runs:
                        run.font.size = Pt(7)
                        run.bold = False

                    # salto de línea extra para separar tareas
                    c_act.add_paragraph("")

        # Horas (mostrar tal cual el texto ingresado)
        horas_text = dia_info.get("horas", "")
        p_horas = c_horas.paragraphs[0]
        p_horas.text = horas_text
        p_horas.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_horas.runs:
            run.font.size = Pt(9)
            run.bold = False
        c_horas.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Acumular minutos
        if dia_info.get("laborable", False) and horas_text:
            h_text = _normalize_dash(horas_text)
            if "–" in h_text:
                total_minutos += parse_hora_rango_a_minutos(h_text)
            else:
                # formato '7H 30M' o '7H' o '90M'
                try:
                    h = 0
                    m = 0
                    for part in h_text.upper().split():
                        part = part.strip()
                        if part.endswith("H"):
                            h = int(part.replace("H", "").strip())
                        elif part.endswith("M"):
                            m = int(part.replace("M", "").strip())
                        elif ":" in part:
                            # posiblemente sola hora '07:30' -> ignoramos (no tiene rango)
                            pass
                    total_minutos += h * 60 + m
                except Exception:
                    pass

        set_row_height(tabla.rows[i], height_cm=data_row_height_cm, rule="atLeast")

    # Fila total
    fila_total = tabla.rows[-1].cells
    fila_total[0].text = ""
    fila_total[1].text = "TOTAL"
    p_total_label = fila_total[1].paragraphs[0]
    p_total_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_total_label.runs:
        run.font.size = Pt(9)
        run.bold = True

    total_str = minutos_a_horas_minutos_str(total_minutos)
    p_total = fila_total[2].paragraphs[0]
    p_total.text = total_str
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_total.runs:
        run.font.size = Pt(9)
        run.bold = True
    fila_total[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(tabla.rows[-1], height_cm=1.0, rule="atLeast")

    return tabla


# --- Buscar e insertar en ancla ---
def find_paragraph_index_with_anchor(doc, anchor):
    for idx, para in enumerate(doc.paragraphs):
        if anchor in para.text:
            return idx
    return None


def insert_table_after_paragraph(doc, paragraph_index, dias_semana, **table_kwargs):
    tabla = build_table_in_doc(doc, dias_semana, **table_kwargs)
    if paragraph_index is None:
        return tabla
    tbl_elm = tabla._tbl
    target_para = doc.paragraphs[paragraph_index]
    target_para._p.addnext(tbl_elm)
    return tabla


# --- Función principal para Flask ---
def generar_cuadro(archivo_base, archivo_salida, dias_semana, ancla="[[AQUI_TABLA]]"):
    """
    - archivo_base: ruta a .docx base
    - archivo_salida: ruta de salida (puede ser un path temporal)
    - dias_semana: lista con diccionarios como en tu ejemplo
    - ancla: texto donde insertar la tabla
    Devuelve Path(archivo_salida)
    """
    doc = Document(archivo_base)
    idx = find_paragraph_index_with_anchor(doc, ancla)
    insert_table_after_paragraph(doc, idx, dias_semana, day_w=1.10, hours_w=1.10, data_row_height_cm=2.8)
    doc.save(archivo_salida)
    return Path(archivo_salida)
